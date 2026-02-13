"""Generate the Agent Edge technical document for AI Futures Lab."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles[f"Heading {level}"]
    heading_style.font.name = "Calibri"
    heading_style.font.color.rgb = RGBColor(0x0A, 0x0E, 0x17)
    if level == 1:
        heading_style.font.size = Pt(24)
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading_style.font.size = Pt(18)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)
    else:
        heading_style.font.size = Pt(14)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

# Helper functions
def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x24, 0x38)

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f" {text}")
    else:
        p.add_run(text)

def add_link_paragraph(doc, label, url):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(10)
    link_run = p.add_run(url)
    link_run.font.size = Pt(10)
    link_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    link_run.underline = True

DIAGRAMS = os.path.join(os.path.dirname(__file__), "diagrams")

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AGENT EDGE")
run.font.size = Pt(36)
run.font.name = "Calibri"
run.bold = True
run.font.color.rgb = RGBColor(0x0A, 0x0E, 0x17)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Multi-Site Edge Intelligence System")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

doc.add_paragraph()

detail = doc.add_paragraph()
detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = detail.add_run("Technical Documentation for AI Futures Lab")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run("Pedro Falcao Costa")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("February 2026")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
run = tagline.add_run("Pure Edge  |  No Cloud  |  A2A + MCP Protocols  |  LFM 2.5 Reasoning")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
run.italic = True

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (placeholder)
# ============================================================
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. The Vision: Edge Agentic AI",
    "   1.1 Agents Are the New APIs",
    "   1.2 The A2A Protocol: Interoperability at Scale",
    "   1.3 AI at the Edge: Beyond the Cloud",
    "2. System Architecture",
    "   2.1 Architecture Overview",
    "   2.2 Data Flow",
    "   2.3 Anomaly Detection Sequence",
    "3. Hardware & Sensor Module",
    "   3.1 ENS160+AHT21 Sensor Module",
    "   3.2 Arduino Integration",
    "4. Protocol Stack",
    "   4.1 MCP: Agent-to-System Integration",
    "   4.2 A2A: Agent-to-Agent Collaboration",
    "   4.3 Message Types & Schemas",
    "5. Agent Design",
    "   5.1 Jetson Agent (Site A)",
    "   5.2 Mac Mini Agent (Control Center)",
    "   5.3 LFM 2.5 Integration",
    "6. Small Language Models & On-Device Intelligence",
    "   6.1 The Case for Edge Intelligence",
    "   6.2 Liquid AI: Leading the SLM Revolution",
    "   6.3 Liquid Foundation Models (LFM)",
    "   6.4 Capgemini and Liquid AI Partnership",
    "   6.5 Why LFM for Agent Edge",
    "7. Dashboard",
    "8. Anomaly Detection & Thresholds",
    "9. Deployment",
    "10. Demo Scenario",
    "11. References & Further Reading",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0]
    run.font.size = Pt(11)
    if not item.startswith("   "):
        run.bold = True

doc.add_page_break()

# ============================================================
# CHAPTER 1: THE VISION
# ============================================================
doc.add_heading("1. The Vision: Edge Agentic AI", level=1)

doc.add_paragraph(
    "We are witnessing a fundamental shift in how distributed systems are designed and operated. "
    "The era of passive microservices responding to predetermined API calls is giving way to a new paradigm: "
    "autonomous, intelligent agents that perceive their environment, reason about it, and collaborate with "
    "peers to make decisions\u2014all without cloud dependency."
)
doc.add_paragraph(
    "Agent Edge is a demonstration of this vision. It brings together three converging trends that are "
    "reshaping enterprise and industrial AI:"
)
add_bullet(doc, "Agent-Oriented Architectures replacing traditional SOA and microservices", bold_prefix="Agents as the new APIs:")
add_bullet(doc, "Google\u2019s A2A protocol enabling cross-vendor, cross-device agent collaboration", bold_prefix="Interoperability protocols:")
add_bullet(doc, "Small, powerful language models (LFM2.5-1.2B-Thinking) running inference locally on edge hardware", bold_prefix="Edge-native intelligence:")

doc.add_paragraph(
    "This system is not a prototype of the future. It is a working demonstration of what is possible today, "
    "using off-the-shelf hardware and open protocols."
)

# --- 1.1 Agents Are the New APIs ---
doc.add_heading("1.1 Agents Are the New APIs", level=2)

doc.add_paragraph(
    "Traditional Service-Oriented Architectures defined rigid contracts between services: fixed endpoints, "
    "predetermined request/response schemas, brittle integrations that break when any component changes. "
    "Agent-Oriented Architectures introduce intelligent, autonomous agents capable of proactively perceiving, "
    "adapting, and dynamically interacting with their environment."
)
doc.add_paragraph(
    "In this new paradigm, agents serve as intelligent intermediaries that:"
)
add_bullet(doc, "Act as secure gatekeepers between disparate systems")
add_bullet(doc, "Breathe new life into legacy infrastructure without costly replacements")
add_bullet(doc, "Enable seamless communication between diverse components through natural language and structured protocols")
add_bullet(doc, "Maintain system integrity while fostering innovation")

doc.add_paragraph(
    "The implications for IoT and Industry 4.0 are profound. We are moving from the Internet of Things\u2014"
    "passive sensor networks that merely transmit data\u2014to the Internet of Agents: dynamic systems where "
    "entities actively communicate, negotiate, and coordinate. Industrial devices adapt autonomously to "
    "operational conditions. Autonomous vehicle fleets negotiate routes in real time. Power grids self-balance "
    "supply and demand. Smart cities become truly intelligent."
)

doc.add_paragraph(
    "Agent Edge demonstrates this concretely: an environmental sensor does not just report numbers. "
    "An agent reasons about those numbers, detects anomalies, and collaborates with a peer agent to "
    "reach a decision\u2014all at the edge, in real time."
)

add_link_paragraph(doc,
    "Full article",
    "https://www.linkedin.com/pulse/agents-new-apis-pedro-falc%C3%A3o-costa-1ydpf/")

# --- 1.2 A2A Protocol ---
doc.add_heading("1.2 The A2A Protocol: Interoperability at Scale", level=2)

doc.add_paragraph(
    "Google\u2019s Agent-to-Agent (A2A) protocol is a major step toward the distributed, modular, and "
    "collaborative AI architectures that the industry needs. For the first time, AI agents from different "
    "vendors, clouds, and devices can communicate using a shared language."
)
doc.add_paragraph(
    "A2A solves the fundamental isolation problem: agents trapped in proprietary runtimes can finally work "
    "together across organizational boundaries. The protocol provides:"
)
add_bullet(doc, "Standardized message types for discovery, observation, analysis, and decision-making")
add_bullet(doc, "Built-in security safeguards for collaborative agent networks")
add_bullet(doc, "Scalability from single-agent to multi-agent ecosystems")
add_bullet(doc, "Edge AI integration for IoT devices and isolated systems")

doc.add_paragraph(
    "Consider the practical implications: an AI-enabled coffee machine detects performance issues and "
    "autonomously contacts maintenance agents, triggering scheduling and parts ordering without human "
    "intervention. Solar farm agents monitoring panel output, battery storage, and weather forecasting "
    "collaborate in real time to optimize energy management."
)
doc.add_paragraph(
    "Agent Edge implements A2A as its core inter-agent communication layer, demonstrating that this protocol "
    "works today on real hardware at the edge."
)

add_link_paragraph(doc,
    "Full article",
    "https://www.linkedin.com/pulse/a2a-protocol-google-here-changes-everything-pedro-falc%C3%A3o-costa-fggff/")

# --- 1.3 AI at the Edge ---
doc.add_heading("1.3 AI at the Edge: Beyond the Cloud", level=2)

doc.add_paragraph(
    "The release of compact, high-quality language models has fundamentally changed what is possible at the edge. "
    "Models like LFM2.5-1.2B-Thinking can run entirely on local hardware\u2014a Jetson Orin Nano, a Mac Mini\u2014"
    "with no cloud dependency, no data leaving the premises, and no per-token API costs."
)
doc.add_paragraph(
    "This opens transformative possibilities:"
)
add_bullet(doc, "How can AI function in minimal-infrastructure environments?", bold_prefix="Extreme environments:")
add_bullet(doc, "Models operating entirely offline with no cloud data transmission", bold_prefix="Privacy and security:")
add_bullet(doc, "Eliminating per-query cloud costs for continuous monitoring", bold_prefix="Cost reduction:")
add_bullet(doc, "AI reasoning at the point of data generation, not in a distant data center", bold_prefix="Latency elimination:")

doc.add_paragraph(
    "Innovation isn\u2019t just about technology. It\u2019s about reimagining what\u2019s possible. "
    "Agent Edge extends AI capability to scenarios where traditional cloud infrastructure is impractical: "
    "factory floors, remote facilities, mobile platforms, and anywhere that demands real-time local intelligence."
)

add_link_paragraph(doc,
    "Original post",
    "https://www.linkedin.com/posts/pedrofalcaocosta_ai-innovation-edgecomputing-activity-7250913453186441216-vAiO")

doc.add_page_break()

# ============================================================
# CHAPTER 2: ARCHITECTURE
# ============================================================
doc.add_heading("2. System Architecture", level=1)

doc.add_paragraph(
    "Agent Edge is a distributed system with two autonomous agents coordinating via A2A protocol. "
    "The Jetson Orin Nano (Site A) reads environmental sensors through an Arduino via MCP protocol "
    "and performs anomaly detection. The Mac Mini M2 (Control Center) provides historical analysis "
    "and hosts the real-time dashboard. Both agents run LFM2.5-1.2B-Thinking locally for reasoning."
)

# --- 2.1 Architecture Overview ---
doc.add_heading("2.1 Architecture Overview", level=2)

arch_path = os.path.join(DIAGRAMS, "architecture.png")
if os.path.exists(arch_path):
    doc.add_picture(arch_path, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Figure 1: Agent Edge System Architecture")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph(
    "The system uses two complementary protocols:"
)
add_bullet(doc, "Agent-to-System integration. The Jetson agent communicates with the Arduino sensor "
    "hardware through an MCP server that exposes a read_sensor tool via stdio transport.", bold_prefix="MCP (Model Context Protocol):")
add_bullet(doc, "Agent-to-Agent collaboration. The Jetson and Mac Mini agents exchange structured "
    "messages over HTTP POST endpoints, with real-time streaming via WebSocket.", bold_prefix="A2A (Agent-to-Agent):")

# --- 2.2 Data Flow ---
doc.add_heading("2.2 Data Flow", level=2)

flow_path = os.path.join(DIAGRAMS, "data_flow.png")
if os.path.exists(flow_path):
    doc.add_picture(flow_path, width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Figure 2: Sensor Data Flow Through All Layers")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph(
    "Sensor data flows as a consistent JSON schema through every layer of the system. The ENS160+AHT21 "
    "module produces five fields\u2014temperature, humidity, eCO2, TVOC, and AQI\u2014which are preserved "
    "from the Arduino serial output through the MCP server, A2A messages, and into the dashboard visualization."
)

# --- 2.3 Anomaly Sequence ---
doc.add_heading("2.3 Anomaly Detection Sequence", level=2)

seq_path = os.path.join(DIAGRAMS, "anomaly_sequence.png")
if os.path.exists(seq_path):
    doc.add_picture(seq_path, width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Figure 3: Anomaly Detection and Collaborative Analysis Sequence")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph(
    "The sequence diagram above illustrates both normal operation (top) and the full anomaly collaboration "
    "flow (bottom). When an anomaly is detected\u2014such as eCO2 exceeding 1000ppm\u2014the Jetson agent "
    "runs local LFM analysis, then requests historical context from the Mac Mini agent. Both agents\u2019 "
    "LFM reasoning is streamed to the dashboard in real time."
)

doc.add_page_break()

# ============================================================
# CHAPTER 3: HARDWARE & SENSOR
# ============================================================
doc.add_heading("3. Hardware & Sensor Module", level=1)

# --- 3.1 Sensor ---
doc.add_heading("3.1 ENS160+AHT21 Sensor Module", level=2)

doc.add_paragraph(
    "The system uses an ENS160+AHT21 combo module connected to an Arduino via I2C (SDA/SCL). "
    "This module provides five environmental parameters:"
)

# Sensor table
table = doc.add_table(rows=6, cols=4)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Parameter", "Sensor", "Unit", "Range"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

data = [
    ("Temperature", "AHT21", "\u00b0C", "-40 to +85"),
    ("Humidity", "AHT21", "% RH", "0 to 100"),
    ("eCO2", "ENS160", "ppm", "400 to 65000"),
    ("TVOC", "ENS160", "ppb", "0 to 65000"),
    ("AQI", "ENS160", "1\u20135 scale", "1 (Good) to 5 (Unhealthy)"),
]
for row_idx, (param, sensor, unit, range_val) in enumerate(data, 1):
    table.rows[row_idx].cells[0].text = param
    table.rows[row_idx].cells[1].text = sensor
    table.rows[row_idx].cells[2].text = unit
    table.rows[row_idx].cells[3].text = range_val
    for cell in table.rows[row_idx].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph(
    "The ENS160 requires approximately 1 minute of warm-up time after power-on, and approximately "
    "1 hour for full calibration. During warm-up, readings may be unreliable."
)

# --- 3.2 Arduino ---
doc.add_heading("3.2 Arduino Integration", level=2)

doc.add_paragraph("The Arduino reads both sensors via I2C and outputs JSON over serial at 9600 baud every 2 seconds:")

add_code_block(doc, '{"temp":24.5,"humidity":65.2,"eco2":450,"tvoc":120,"aqi":1}')

doc.add_paragraph("Required Arduino libraries:")
add_bullet(doc, "ScioSense_ENS160 \u2014 eCO2, TVOC, AQI readings via I2C")
add_bullet(doc, "Adafruit_AHTX0 \u2014 Temperature and humidity readings via I2C")
add_bullet(doc, "Wire.h \u2014 I2C communication")

doc.add_paragraph(
    "The AHT21 temperature and humidity readings are fed to the ENS160 for environmental compensation, "
    "improving the accuracy of the air quality measurements."
)

doc.add_page_break()

# ============================================================
# CHAPTER 4: PROTOCOL STACK
# ============================================================
doc.add_heading("4. Protocol Stack", level=1)

# --- 4.1 MCP ---
doc.add_heading("4.1 MCP: Agent-to-System Integration", level=2)

doc.add_paragraph(
    "The Model Context Protocol (MCP) provides a standardized way for AI agents to interact with external "
    "tools and data sources. In Agent Edge, the MCP Arduino Server exposes the sensor hardware as a tool "
    "that the Jetson agent can invoke."
)
add_bullet(doc, "read_sensor \u2014 returns current sensor readings with timestamp", bold_prefix="MCP Tool:")
add_bullet(doc, "stdio (subprocess communication)", bold_prefix="Transport:")
add_bullet(doc, "JSON with schema: temperature, humidity, eco2, tvoc, aqi, timestamp", bold_prefix="Return format:")

# --- 4.2 A2A ---
doc.add_heading("4.2 A2A: Agent-to-Agent Collaboration", level=2)

doc.add_paragraph(
    "The A2A protocol enables peer-to-peer agent communication. Each agent exposes:"
)
add_bullet(doc, "/a2a/message \u2014 HTTP POST endpoint for receiving structured messages")
add_bullet(doc, "/stream \u2014 WebSocket endpoint for real-time event streaming")
add_bullet(doc, "/health \u2014 Agent Card and status endpoint")

# --- 4.3 Message Types ---
doc.add_heading("4.3 Message Types & Schemas", level=2)

msg_table = doc.add_table(rows=7, cols=3)
msg_table.style = "Light Grid Accent 1"
msg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
msg_headers = ["Message Type", "Direction", "Purpose"]
for i, h in enumerate(msg_headers):
    cell = msg_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

msg_data = [
    ("agent_card", "Broadcast", "Agent discovery and capability advertisement"),
    ("sensor_observation", "Jetson \u2192 Mac Mini", "Sensor reading with all 5 fields + location"),
    ("analysis_request", "Jetson \u2192 Mac Mini", "Anomaly detected, requesting historical analysis"),
    ("analysis_response", "Mac Mini \u2192 Jetson", "Historical context and LFM reasoning"),
    ("decision", "Both agents", "Collaborative decision logged by both"),
    ("heartbeat", "Bidirectional", "Periodic health check (every 10s)"),
]
for row_idx, (mtype, direction, purpose) in enumerate(msg_data, 1):
    msg_table.rows[row_idx].cells[0].text = mtype
    msg_table.rows[row_idx].cells[1].text = direction
    msg_table.rows[row_idx].cells[2].text = purpose
    for cell in msg_table.rows[row_idx].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

doc.add_paragraph("Sensor observation payload schema:")
add_code_block(doc,
    '{\n'
    '  "type": "sensor_observation",\n'
    '  "from": "jetson-site-a",\n'
    '  "to": "macmini-control",\n'
    '  "message_id": "uuid",\n'
    '  "timestamp": "ISO8601",\n'
    '  "payload": {\n'
    '    "sensor": "ENS160+AHT21",\n'
    '    "temperature": 24.5,\n'
    '    "humidity": 65.2,\n'
    '    "eco2": 450,        // ppm\n'
    '    "tvoc": 120,        // ppb\n'
    '    "aqi": 1,           // 1-5 scale\n'
    '    "location": "Site A - Server Room"\n'
    '  }\n'
    '}'
)

doc.add_page_break()

# ============================================================
# CHAPTER 5: AGENT DESIGN
# ============================================================
doc.add_heading("5. Agent Design", level=1)

# --- 5.1 Jetson ---
doc.add_heading("5.1 Jetson Agent (Site A)", level=2)

doc.add_paragraph(
    "The Jetson Orin Nano runs the Site A monitoring agent. It is responsible for sensor reading, "
    "anomaly detection, and initiating collaborative analysis."
)
doc.add_paragraph("Responsibilities:")
add_bullet(doc, "Initialize MCP client to Arduino sensor server", bold_prefix="Startup:")
add_bullet(doc, "Read sensor every 5 seconds via MCP, send observations to Mac Mini", bold_prefix="Sensor loop:")
add_bullet(doc, "Check 4 thresholds on every reading; trigger LFM analysis if any exceeded", bold_prefix="Anomaly detection:")
add_bullet(doc, "Process analysis responses, log collaborative decisions", bold_prefix="A2A handling:")

doc.add_paragraph("Configuration:")
add_code_block(doc,
    'AGENT_ID = "jetson-site-a"\n'
    'AGENT_PORT = 8080\n'
    'SENSOR_POLL_INTERVAL = 5  # seconds\n'
    'TEMP_DELTA_THRESHOLD = 5.0  # degrees C\n'
    'ECO2_THRESHOLD = 1000  # ppm\n'
    'TVOC_THRESHOLD = 500  # ppb\n'
    'AQI_THRESHOLD = 4  # 1-5 scale'
)

# --- 5.2 Mac Mini ---
doc.add_heading("5.2 Mac Mini Agent (Control Center)", level=2)

doc.add_paragraph(
    "The Mac Mini M2 runs the Control Center agent. It maintains historical sensor data, performs "
    "statistical analysis, and provides context for anomaly evaluation."
)
doc.add_paragraph("Responsibilities:")
add_bullet(doc, "Maintain rolling window of sensor readings (all 5 fields, last 24 hours)")
add_bullet(doc, "Calculate per-field statistics: mean, standard deviation, min, max")
add_bullet(doc, "Provide multi-parameter context (e.g., \u201cCO2 spiking while temperature stable\u201d)")
add_bullet(doc, "Run LFM analysis with historical context when analysis requests arrive")
add_bullet(doc, "Host dashboard data APIs and WebSocket streams")

# --- 5.3 LFM ---
doc.add_heading("5.3 LFM 2.5 Integration", level=2)

doc.add_paragraph(
    "Both agents run Liquid AI\u2019s LFM2.5-1.2B-Thinking model locally. The Jetson uses CUDA acceleration; "
    "the Mac Mini uses Apple\u2019s MPS (Metal Performance Shaders). The LFM client supports:"
)
add_bullet(doc, "Token-by-token streaming for real-time dashboard visualization")
add_bullet(doc, "Structured prompts for anomaly detection, historical analysis, and collaborative reasoning")
add_bullet(doc, "GPU-accelerated inference with automatic device detection (CUDA/MPS/CPU)")

doc.add_page_break()

# ============================================================
# CHAPTER 6: SLMs & ON-DEVICE INTELLIGENCE
# ============================================================
doc.add_heading("6. Small Language Models & On-Device Intelligence", level=1)

doc.add_paragraph(
    "The AI industry has spent the last several years in a relentless pursuit of scale: larger models, "
    "bigger clusters, more parameters. But a parallel revolution has been quietly gaining momentum\u2014one "
    "that may ultimately prove more consequential for how AI is actually deployed in the real world. "
    "Small Language Models (SLMs) are redefining what is possible at the edge, on device, and in "
    "environments where cloud connectivity is a luxury, not a given."
)
doc.add_paragraph(
    "Agent Edge is built on a fundamental conviction: the future of AI is not exclusively in the cloud. "
    "It is at the edge. It is on the device. It is wherever decisions need to be made\u2014fast, private, "
    "and without dependency on a data center thousands of kilometers away."
)

# --- 6.1 The Case for Edge Intelligence ---
doc.add_heading("6.1 The Case for Edge Intelligence", level=2)

doc.add_paragraph(
    "Cloud-based AI has an undeniable role, but it comes with constraints that are unacceptable in many "
    "real-world scenarios:"
)
add_bullet(doc, "Every API call is a round trip. For real-time monitoring, autonomous vehicles, "
    "and industrial control, even 100ms of latency is too much.", bold_prefix="Latency:")
add_bullet(doc, "Sensitive sensor data, patient records, factory telemetry\u2014organizations increasingly "
    "cannot or will not send this data to external servers.", bold_prefix="Privacy:")
add_bullet(doc, "Factory floors, offshore platforms, remote facilities, vehicles in tunnels\u2014"
    "many critical environments have intermittent or no connectivity.", bold_prefix="Connectivity:")
add_bullet(doc, "Continuous inference through cloud APIs creates compounding costs that make "
    "always-on monitoring economically impractical.", bold_prefix="Cost:")
add_bullet(doc, "The energy footprint of routing billions of inference requests through "
    "hyperscale data centers is unsustainable at global scale.", bold_prefix="Sustainability:")

doc.add_paragraph(
    "On-device intelligence eliminates these constraints entirely. When a 1.2B parameter model runs "
    "locally on a Jetson Orin Nano or a Mac Mini, inference is instantaneous, data never leaves the "
    "premises, and the system operates independently of network conditions. This is not a compromise\u2014"
    "for edge use cases, it is superior architecture."
)

# --- 6.2 Liquid AI ---
doc.add_heading("6.2 Liquid AI: Leading the SLM Revolution", level=2)

doc.add_paragraph(
    "Liquid AI, a spin-off from MIT\u2019s Computer Science and Artificial Intelligence Laboratory (CSAIL), "
    "has taken a fundamentally different approach to foundation models. Rather than competing in the "
    "parameter arms race, Liquid AI has focused on building the most efficient models at every scale\u2014"
    "models that are designed from the ground up for on-device deployment."
)
doc.add_paragraph(
    "The core innovation is Liquid Neural Networks (LNNs): a new class of neural architecture that is "
    "more robust, computationally smaller, and less memory-intensive than traditional transformer models. "
    "Where conventional LLMs require massive GPU clusters, Liquid\u2019s models achieve competitive quality "
    "while running on CPUs, NPUs, and embedded processors with minimal memory footprint."
)
doc.add_paragraph(
    "Liquid AI\u2019s decision to release open-weight models was a pivotal strategic move. By making LFM "
    "weights available on Hugging Face, they have enabled an entire ecosystem of developers, researchers, "
    "and enterprises to build on their technology. This openness, combined with native support for "
    "llama.cpp, MLX, vLLM, and ONNX across Apple, AMD, Qualcomm, and NVIDIA hardware, positions "
    "Liquid AI to lead the SLM space in a way that closed-model providers cannot."
)
doc.add_paragraph(
    "I firmly believe that Liquid AI will be a defining force in the Small Language Model space. Their "
    "combination of MIT-grade research, open weights, hardware-optimized architectures, and a clear "
    "focus on the edge makes them uniquely positioned to enable the next generation of on-device "
    "intelligence\u2014from wearables and mobile phones to autonomous vehicles and industrial IoT."
)

# --- 6.3 LFM Models ---
doc.add_heading("6.3 Liquid Foundation Models (LFM)", level=2)

doc.add_paragraph(
    "The LFM2.5 family represents the latest generation of Liquid Foundation Models, with extended "
    "pretraining from 10 trillion to 28 trillion tokens and significantly scaled post-training via "
    "reinforcement learning."
)

doc.add_paragraph("The LFM2.5 model family includes:")

lfm_table = doc.add_table(rows=8, cols=3)
lfm_table.style = "Light Grid Accent 1"
lfm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
lfm_headers = ["Model", "Parameters", "Capability"]
for i, h in enumerate(lfm_headers):
    cell = lfm_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

lfm_data = [
    ("LFM2.5-1.2B-Base", "1.17B", "Pretrained checkpoint for fine-tuning"),
    ("LFM2.5-1.2B-Instruct", "1.17B", "General-purpose instruction following"),
    ("LFM2.5-1.2B-Thinking", "1.17B", "On-device reasoning under 1GB"),
    ("LFM2.5-1.2B-JP", "1.17B", "Japanese language optimization"),
    ("LFM2.5-VL-1.6B", "1.6B", "Vision-language understanding"),
    ("LFM2.5-Audio-1.5B", "1.5B", "End-to-end audio and text generation"),
    ("LFM2-8B-A1B", "8B (1B active)", "Mixture-of-experts efficiency"),
]
for row_idx, (model, params, cap) in enumerate(lfm_data, 1):
    lfm_table.rows[row_idx].cells[0].text = model
    lfm_table.rows[row_idx].cells[1].text = params
    lfm_table.rows[row_idx].cells[2].text = cap
    for cell in lfm_table.rows[row_idx].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

doc.add_paragraph("Key technical characteristics:")
add_bullet(doc, "16-layer hybrid architecture: 10 double-gated LIV convolution blocks + 6 GQA attention blocks",
    bold_prefix="Architecture:")
add_bullet(doc, "32,768 tokens, enabling complex RAG tasks and long-form data extraction",
    bold_prefix="Context length:")
add_bullet(doc, "2,975 tok/s prefill, 116 tok/s decode on AMD Ryzen AI 9; 335 tok/s prefill on Snapdragon Gen4 mobile CPU",
    bold_prefix="Inference speed:")
add_bullet(doc, "English, Arabic, Chinese, French, German, Japanese, Korean, Spanish",
    bold_prefix="Languages:")
add_bullet(doc, "Vehicles, mobile devices, laptops, IoT devices, embedded systems, edge servers",
    bold_prefix="Deployment targets:")

doc.add_paragraph(
    "The LFM2.5-1.2B-Thinking variant is particularly relevant for Agent Edge: it enables on-device "
    "reasoning in under 1GB of memory, making it feasible to run chain-of-thought analysis on edge "
    "hardware for anomaly detection and collaborative decision-making."
)

# --- 6.4 Capgemini Partnership ---
doc.add_heading("6.4 Capgemini and Liquid AI Partnership", level=2)

doc.add_paragraph(
    "Capgemini and Liquid AI have established a strategic collaboration to build next-generation AI "
    "solutions for enterprises. Capgemini\u2019s Corporate VC fund (ISAI Cap Venture) participated in "
    "Liquid AI\u2019s $37.6 million seed round in 2023, and the partnership has since deepened into a "
    "joint effort to bring on-device AI to production at enterprise scale."
)

doc.add_paragraph(
    "The collaboration focuses on deploying smaller, more sustainable generative AI models compatible "
    "with offline and edge devices, with applications spanning manufacturing, healthcare, finance, "
    "and industrial IoT. As Andrew Vickers, CTO Generative AI at Capgemini Engineering, stated:"
)

quote = doc.add_paragraph()
quote.paragraph_format.left_indent = Cm(1.5)
quote.paragraph_format.right_indent = Cm(1.5)
run = quote.add_run(
    "\u201cWe\u2019re thrilled to be continuing our relationship with Liquid AI, a relationship born "
    "from our Strategic University Research Partnership Program with MIT. Liquid AI brings amazing "
    "technology to the market, in tune with the needed trends of simplicity, efficiency and sustainability.\u201d"
)
run.italic = True
run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

doc.add_paragraph(
    "Ramin Hasani, Co-Founder and CEO of Liquid AI, added:"
)

quote2 = doc.add_paragraph()
quote2.paragraph_format.left_indent = Cm(1.5)
quote2.paragraph_format.right_indent = Cm(1.5)
run2 = quote2.add_run(
    "\u201cOur cooperation with Capgemini allows us to combine our cutting-edge AI technology with "
    "Capgemini\u2019s global reach and expertise, ultimately benefiting organizations across various "
    "industries to develop and deploy private, reliable and best-in-class domain-specific and "
    "generalist AI systems.\u201d"
)
run2.italic = True
run2.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

doc.add_paragraph(
    "Agent Edge is a direct expression of this partnership\u2019s vision: Liquid AI\u2019s models running "
    "on edge hardware, orchestrated by Capgemini\u2019s expertise in distributed systems and agent "
    "architectures, demonstrated through Capgemini\u2019s AI Futures Lab (Generative AI Lab)."
)

add_link_paragraph(doc,
    "Capgemini AI Futures (Generative AI Lab)",
    "https://www.capgemini.com/about-us/who-we-are/innovation-ecosystem/aifutures/")

add_link_paragraph(doc,
    "Capgemini and Liquid AI Partnership",
    "https://www.capgemini.com/capgemini-and-liquid-ai/")

add_link_paragraph(doc,
    "Liquid AI Collaboration Announcement",
    "https://www.liquid.ai/blog/liquid-ai-announces-collaboration-with-capgemini-to-build-next-generation-ai-solutions-for-enterprises")

# --- 6.5 Why LFM for Agent Edge ---
doc.add_heading("6.5 Why LFM for Agent Edge", level=2)

doc.add_paragraph(
    "Agent Edge uses the LFM2.5-1.2B model on both agents for a set of deliberate reasons:"
)
add_bullet(doc, "At 1.17B parameters, the model fits comfortably on the Jetson Orin Nano (CUDA) "
    "and Mac Mini M2 (MPS) with memory to spare for the agent runtime.",
    bold_prefix="Right-sized for edge hardware:")
add_bullet(doc, "Token-by-token streaming enables the dashboard to show AI reasoning in real time\u2014"
    "the key differentiator for the demo.",
    bold_prefix="Streaming inference:")
add_bullet(doc, "The model handles anomaly detection prompts, historical analysis, and collaborative "
    "reasoning with sufficient quality for the monitoring use case.",
    bold_prefix="Quality at scale:")
add_bullet(doc, "No cloud dependency means the demo works anywhere\u2014conference halls, factory floors, "
    "air-gapped environments.",
    bold_prefix="True edge deployment:")
add_bullet(doc, "Open weights on Hugging Face mean the audience can reproduce and extend the demo "
    "with their own hardware.",
    bold_prefix="Open and reproducible:")

doc.add_paragraph(
    "This is the promise of on-device intelligence made tangible: two agents, on two different hardware "
    "platforms, running local language models, collaborating through open protocols, reasoning about "
    "real sensor data\u2014all without a single byte leaving the room."
)

doc.add_page_break()

# ============================================================
# CHAPTER 7: DASHBOARD
# ============================================================
doc.add_heading("7. Dashboard", level=1)

doc.add_paragraph(
    "The Streamlit dashboard provides a real-time \u201cEdge Operations Center\u201d view of the entire system. "
    "It connects to both agents via HTTP and WebSocket, displaying sensor data, agent status, A2A message flow, "
    "and LFM reasoning streams."
)

doc.add_paragraph("Dashboard components:")
add_bullet(doc, "Live connection indicators with Agent Card details for both agents", bold_prefix="Agent Status:")
add_bullet(doc, "Real-time Plotly charts for Temperature+Humidity and eCO2+TVOC, "
    "plus an AQI gauge indicator (1\u20135 scale) with anomaly markers", bold_prefix="Sensor Visualization:")
add_bullet(doc, "Scrolling, color-coded message feed showing all inter-agent communication", bold_prefix="A2A Message Flow:")
add_bullet(doc, "Token-by-token display of LFM analysis with anomaly trigger context "
    "and cross-parameter reasoning", bold_prefix="LFM Reasoning Stream:")

doc.add_paragraph(
    "The dashboard aesthetic is \u201cindustrial precision meets data-viz clarity\u201d\u2014a dark theme with "
    "electric cyan for live data, amber for warnings, and crimson for alerts. Typography uses JetBrains Mono "
    "for data and Outfit for headers."
)

doc.add_page_break()

# ============================================================
# CHAPTER 8: ANOMALY DETECTION
# ============================================================
doc.add_heading("8. Anomaly Detection & Thresholds", level=1)

doc.add_paragraph(
    "The Jetson agent evaluates every sensor reading against four independent thresholds. If any single "
    "threshold is exceeded, the full anomaly analysis pipeline is triggered."
)

threshold_table = doc.add_table(rows=5, cols=4)
threshold_table.style = "Light Grid Accent 1"
threshold_table.alignment = WD_TABLE_ALIGNMENT.CENTER
th_headers = ["Parameter", "Threshold", "Condition", "Meaning"]
for i, h in enumerate(th_headers):
    cell = threshold_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

th_data = [
    ("Temperature", "> 5\u00b0C delta", "vs. previous reading", "Rapid temperature change"),
    ("eCO2", "> 1000 ppm", "absolute value", "Poor indoor air quality"),
    ("TVOC", "> 500 ppb", "absolute value", "Elevated volatile organic compounds"),
    ("AQI", "\u2265 4", "absolute value", "Unhealthy air quality (ENS160 scale)"),
]
for row_idx, (param, threshold, condition, meaning) in enumerate(th_data, 1):
    threshold_table.rows[row_idx].cells[0].text = param
    threshold_table.rows[row_idx].cells[1].text = threshold
    threshold_table.rows[row_idx].cells[2].text = condition
    threshold_table.rows[row_idx].cells[3].text = meaning
    for cell in threshold_table.rows[row_idx].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

doc.add_paragraph("When triggered, the anomaly pipeline:")
add_bullet(doc, "Jetson runs local LFM analysis with current + previous sensor readings")
add_bullet(doc, "Jetson sends analysis_request to Mac Mini with anomaly context and LFM reasoning")
add_bullet(doc, "Mac Mini queries 24-hour historical statistics for all sensor fields")
add_bullet(doc, "Mac Mini runs LFM analysis with historical context")
add_bullet(doc, "Mac Mini returns analysis_response with confidence score")
add_bullet(doc, "Both agents log the collaborative decision")
add_bullet(doc, "All reasoning is streamed to the dashboard in real time")

doc.add_page_break()

# ============================================================
# CHAPTER 9: DEPLOYMENT
# ============================================================
doc.add_heading("9. Deployment", level=1)

doc.add_heading("Project Structure", level=2)
add_code_block(doc,
    'agent-edge/\n'
    '\u251c\u2500\u2500 arduino/sensor_reader/sensor_reader.ino\n'
    '\u251c\u2500\u2500 mcp_servers/arduino/server.py\n'
    '\u251c\u2500\u2500 shared/\n'
    '\u2502   \u251c\u2500\u2500 a2a_protocol.py\n'
    '\u2502   \u251c\u2500\u2500 agent_card.py\n'
    '\u2502   \u251c\u2500\u2500 lfm_client.py\n'
    '\u2502   \u2514\u2500\u2500 storage.py\n'
    '\u251c\u2500\u2500 agents/\n'
    '\u2502   \u251c\u2500\u2500 jetson/  (agent.py, config.py, mcp_client.py)\n'
    '\u2502   \u2514\u2500\u2500 macmini/ (agent.py, config.py)\n'
    '\u251c\u2500\u2500 dashboard/\n'
    '\u2502   \u251c\u2500\u2500 app.py\n'
    '\u2502   \u2514\u2500\u2500 components/ (sensor_viz, a2a_conversation, lfm_reasoning, agent_status)\n'
    '\u251c\u2500\u2500 scripts/ (setup_jetson.sh, setup_macmini.sh, run_demo.sh)\n'
    '\u2514\u2500\u2500 tests/'
)

doc.add_heading("Quick Start", level=2)

doc.add_paragraph("Jetson Orin Nano (Site A):")
add_code_block(doc,
    'bash scripts/setup_jetson.sh\n'
    'export MACMINI_AGENT_URL=http://<macmini-ip>:8081\n'
    'export SERIAL_PORT=/dev/ttyACM0\n'
    'python -m agents.jetson.agent'
)

doc.add_paragraph("Mac Mini M2 (Control Center):")
add_code_block(doc,
    'bash scripts/setup_macmini.sh\n'
    'export JETSON_AGENT_URL=http://<jetson-ip>:8080\n'
    'python -m agents.macmini.agent\n'
    'streamlit run dashboard/app.py'
)

doc.add_heading("Technology Stack", level=2)

tech_table = doc.add_table(rows=8, cols=2)
tech_table.style = "Light Grid Accent 1"
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
tech_table.rows[0].cells[0].text = "Component"
tech_table.rows[0].cells[1].text = "Technology"
for paragraph in tech_table.rows[0].cells[0].paragraphs:
    for run in paragraph.runs:
        run.bold = True
for paragraph in tech_table.rows[0].cells[1].paragraphs:
    for run in paragraph.runs:
        run.bold = True

tech_data = [
    ("Language", "Python 3.10+"),
    ("Web Framework", "FastAPI + Uvicorn"),
    ("Agent Communication", "A2A Protocol (HTTP POST + WebSocket)"),
    ("Tool Integration", "MCP (Model Context Protocol)"),
    ("LLM", "Liquid AI LFM2.5-1.2B-Thinking (HuggingFace Transformers)"),
    ("Dashboard", "Streamlit + Plotly"),
    ("Storage", "JSON Lines (append-only, no database)"),
]
for row_idx, (comp, tech) in enumerate(tech_data, 1):
    tech_table.rows[row_idx].cells[0].text = comp
    tech_table.rows[row_idx].cells[1].text = tech

doc.add_page_break()

# ============================================================
# CHAPTER 10: DEMO SCENARIO
# ============================================================
doc.add_heading("10. Demo Scenario", level=1)

doc.add_paragraph(
    "The following scenario demonstrates the full Agent Edge system in action:"
)

doc.add_heading("Act 1: System Online", level=3)
doc.add_paragraph(
    "Both agents start and discover each other. The dashboard shows green status indicators. "
    "Sensor data begins flowing: temperature around 24\u00b0C, humidity 65%, eCO2 around 450ppm, "
    "AQI 1 (Good). The audience sees real data on real hardware."
)

doc.add_heading("Act 2: Normal Monitoring", level=3)
doc.add_paragraph(
    "Every 5 seconds, the Jetson reads the sensor via MCP and sends an A2A observation to the Mac Mini. "
    "The dashboard charts fill with data. The A2A message feed scrolls. Everything is within normal thresholds."
)

doc.add_heading("Act 3: Anomaly Triggers", level=3)
doc.add_paragraph(
    "CO2 levels begin rising (simulated by breathing near the sensor or using a CO2 source). "
    "When eCO2 crosses 1000ppm, the Jetson\u2019s anomaly detector fires. The dashboard turns amber."
)

doc.add_heading("Act 4: Collaborative Intelligence", level=3)
doc.add_paragraph(
    "The Jetson runs LFM analysis locally: \u201cCO2 spike while temperature is stable suggests poor "
    "ventilation rather than a heat source.\u201d It sends an analysis request to the Mac Mini. "
    "The Mac Mini queries its 24-hour history, computes statistics, and runs its own LFM analysis: "
    "\u201ceCO2 is 2.1 standard deviations above the historical mean. Recommend checking ventilation.\u201d"
)
doc.add_paragraph(
    "The audience watches both LFM reasoning streams appear token-by-token on the dashboard. "
    "Two agents, on two different hardware platforms, collaborating to reach a decision\u2014"
    "all at the edge, no cloud involved."
)

doc.add_heading("Act 5: Resolution", level=3)
doc.add_paragraph(
    "The collaborative decision is logged by both agents. The dashboard shows the decision history. "
    "As CO2 levels return to normal, the system resumes routine monitoring."
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Key talking points for the audience:")
run.bold = True

add_bullet(doc, "Pure edge: zero cloud dependency, zero data leaving the room")
add_bullet(doc, "Open protocols: A2A and MCP are interoperability standards, not proprietary")
add_bullet(doc, "Visible reasoning: the audience can watch AI think, not just see outputs")
add_bullet(doc, "Real hardware: Jetson + Arduino + real sensors, not a simulation")
add_bullet(doc, "Practical: same architecture scales to industrial monitoring, smart buildings, fleet management")

doc.add_page_break()

# ============================================================
# CHAPTER 11: REFERENCES
# ============================================================
doc.add_heading("11. References & Further Reading", level=1)

doc.add_heading("Articles by Pedro Falcao Costa", level=2)

add_link_paragraph(doc,
    "Agents Are the New APIs",
    "https://www.linkedin.com/pulse/agents-new-apis-pedro-falc%C3%A3o-costa-1ydpf/")

add_link_paragraph(doc,
    "The A2A Protocol: Google Is Here and Changes Everything",
    "https://www.linkedin.com/pulse/a2a-protocol-google-here-changes-everything-pedro-falc%C3%A3o-costa-fggff/")

add_link_paragraph(doc,
    "AI, Innovation & Edge Computing",
    "https://www.linkedin.com/posts/pedrofalcaocosta_ai-innovation-edgecomputing-activity-7250913453186441216-vAiO")

doc.add_heading("Protocols & Standards", level=2)

add_bullet(doc, "Model Context Protocol (MCP) \u2014 Anthropic, https://modelcontextprotocol.io")
add_bullet(doc, "Agent-to-Agent Protocol (A2A) \u2014 Google, https://google.github.io/A2A/")

doc.add_heading("Liquid AI & Capgemini", level=2)

add_link_paragraph(doc,
    "Liquid AI",
    "https://www.liquid.ai/")

add_link_paragraph(doc,
    "Liquid Foundation Models",
    "https://www.liquid.ai/models")

add_link_paragraph(doc,
    "LFM2.5 Announcement",
    "https://www.liquid.ai/blog/introducing-lfm2-5-the-next-generation-of-on-device-ai")

add_link_paragraph(doc,
    "Capgemini and Liquid AI Partnership",
    "https://www.capgemini.com/capgemini-and-liquid-ai/")

add_link_paragraph(doc,
    "Capgemini AI Futures (Generative AI Lab)",
    "https://www.capgemini.com/about-us/who-we-are/innovation-ecosystem/aifutures/")

add_link_paragraph(doc,
    "What Are Liquid Neural Networks? (Capgemini)",
    "https://www.capgemini.com/insights/expert-perspectives/what-are-liquid-neural-networks-and-why-should-you-care/")

add_link_paragraph(doc,
    "LFM2.5-1.2B-Instruct on Hugging Face",
    "https://huggingface.co/LiquidAI/LFM2.5-1.2B-Instruct")

doc.add_heading("Hardware & Sensors", level=2)

add_bullet(doc, "ENS160 Digital Multi-Gas Sensor \u2014 ScioSense")
add_bullet(doc, "AHT21 Temperature & Humidity Sensor \u2014 Aosong")
add_bullet(doc, "NVIDIA Jetson Orin Nano \u2014 NVIDIA")

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), "Agent_Edge_AI_Futures_Lab.docx")
doc.save(output_path)
print(f"Document saved: {output_path}")
